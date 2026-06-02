import os
import tempfile
from html import escape
from io import BytesIO
from pathlib import Path

from docx import Document
from pypdf import PdfReader


def _read_pdf(content: bytes) -> PdfReader:
    return PdfReader(BytesIO(content))


def _page_texts(content: bytes) -> list[str]:
    reader = _read_pdf(content)
    return [(page.extract_text() or "") for page in reader.pages]


async def pdf_to_text(content: bytes, filename: str) -> tuple[str, str]:
    text = "\n\n".join(_page_texts(content))
    stem = Path(filename).stem
    out_path = os.path.join(tempfile.mkdtemp(prefix="ff_out_"), f"{stem}.txt")
    Path(out_path).write_text(text, encoding="utf-8")
    return out_path, f"{stem}_converted.txt"


async def pdf_to_docx(content: bytes, filename: str) -> tuple[str, str]:
    document = Document()

    for page_number, text in enumerate(_page_texts(content), start=1):
        if page_number > 1:
            document.add_page_break()
        text = text.strip()
        if text:
            for paragraph in text.split("\n\n"):
                document.add_paragraph(paragraph.replace("\n", " "))
        else:
            document.add_paragraph(f"Page {page_number} did not contain selectable text.")

    stem = Path(filename).stem
    out_path = os.path.join(tempfile.mkdtemp(prefix="ff_out_"), f"{stem}.docx")
    document.save(out_path)
    return out_path, f"{stem}_converted.docx"


async def pdf_to_html(content: bytes, filename: str) -> tuple[str, str]:
    html_parts = ['<!doctype html>', '<html><body style="font-family: sans-serif; line-height: 1.5;">']
    for page_number, text in enumerate(_page_texts(content), start=1):
        if page_number > 1:
            html_parts.append("<hr>")
        html_parts.append(f"<h2>Page {page_number}</h2>")
        escaped_text = escape(text.strip())
        html_parts.append(f"<pre style=\"white-space: pre-wrap;\">{escaped_text}</pre>")
    html_parts.append("</body></html>")

    stem = Path(filename).stem
    out_path = os.path.join(tempfile.mkdtemp(prefix="ff_out_"), f"{stem}.html")
    Path(out_path).write_text("\n".join(html_parts), encoding="utf-8")
    return out_path, f"{stem}_converted.html"
